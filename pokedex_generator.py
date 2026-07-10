# This project is not affiliated with Nintendo, Game Freak, or The Pokémon Company.
# It uses publicly available data from PokéAPI for educational and non-commercial use.

import asyncio
import aiohttp
import os
from docx import Document
from docx.shared import Inches

async def fetch_json(session, url):
    async with session.get(url) as response:
        return await response.json()

async def download_image(session, url, filepath):
    if not url:
        return
    async with session.get(url) as response:
        img_bytes = await response.read()
        with open(filepath, "wb") as f:
            f.write(img_bytes)

async def get_all_pokemon():
    async with aiohttp.ClientSession() as session:
        async with session.get('https://pokeapi.co/api/v2/pokemon?limit=151') as list_response:
            list_data = await list_response.json()

        tasks = [fetch_json(session, pokemon['url']) for pokemon in list_data['results']]
        return await asyncio.gather(*tasks)

def extract_english_flavor_texts(species_json):
    entries = species_json["flavor_text_entries"]

    english_entries = [
        e["flavor_text"].replace("\n", " ").replace("\f", " ")
        for e in entries
        if e["language"]["name"] == "en"
    ]

    # Remove duplicates while preserving order
    seen = set()
    unique_entries = []
    for ft in english_entries:
        if ft not in seen:
            seen.add(ft)
            unique_entries.append(ft)

    return unique_entries

def extract_evolution_chain(evo_json):
    chain = []
    current = evo_json["chain"]

    while current:
        chain.append(current["species"]["name"])
        current = current["evolves_to"][0] if current["evolves_to"] else None

    return chain

def dm_to_height(dm):
    meters = dm / 10
    total_inches = meters * 39.3701
    feet = int(total_inches // 12)
    inches = int(round(total_inches % 12))
    return f"{feet}'{inches:02}\" ({meters:.1f} m)"

def hg_to_weight(hg):
    kg = hg / 10
    lbs = kg * 2.20462
    return f"{lbs:.1f} lbs ({kg:.1f} kg)"

def build_docx(pokemon_id, pokemon_name, summary, front_path, art_path):
    doc = Document()

    # Images only at the top (left aligned)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    if os.path.exists(front_path):
        cell.add_paragraph().add_run().add_picture(front_path, width=Inches(1.5))
    if os.path.exists(art_path):
        cell.add_paragraph().add_run().add_picture(art_path, width=Inches(1.5))

    doc.add_paragraph("")  # spacing
    doc.add_paragraph(summary)

    docx_filename = f"{pokemon_id:03d}_{pokemon_name}.docx"
    doc.save(os.path.join("pokemon_output", docx_filename))

def summarize_pokemon(p, species, evo_chain, flavor_texts):
    pokemon_id = p["id"]
    pokemon_name = p["name"].capitalize()

    abilities = [a["ability"]["name"] for a in p["abilities"]]
    types = [t["type"]["name"] for t in p["types"]]
    stats = {s["stat"]["name"]: s["base_stat"] for s in p["stats"]}
    moves = [m["move"]["name"] for m in p["moves"]]

    color = species["color"]["name"]
    egg_groups = [g["name"] for g in species["egg_groups"]]
    is_baby = species["is_baby"]
    is_legendary = species["is_legendary"]
    is_mythical = species["is_mythical"]

    height_str = dm_to_height(p["height"])
    weight_str = hg_to_weight(p["weight"])

    lines = []
    lines.append(f"ID: {pokemon_id}")
    lines.append(f"Name: {pokemon_name}")
    lines.append(f"Base Experience Yield: {p['base_experience']}")
    lines.append(f"Height: {height_str}")
    lines.append(f"Weight: {weight_str}")
    lines.append(f"Color: {color}")
    lines.append("")

    lines.append("Abilities:")
    for ability in abilities:
        lines.append(f" - {ability}")

    lines.append("")
    lines.append("Types:")
    for t in types:
        lines.append(f" - {t}")

    lines.append("")
    lines.append("Stats:")
    for stat_name, stat_value in stats.items():
        lines.append(f" - {stat_name}: {stat_value}")

    lines.append("")
    lines.append("Species Information:")
    lines.append(f" - Egg Groups: {', '.join(egg_groups)}")
    lines.append(f" - Is Baby: {is_baby}")
    lines.append(f" - Is Legendary: {is_legendary}")
    lines.append(f" - Is Mythical: {is_mythical}")
    lines.append(f" - Evolution Chain: {' → '.join([e.capitalize() for e in evo_chain])}")

    lines.append("")
    lines.append("Flavor Text in the games (English):")
    for ft in flavor_texts:
        lines.append(f" - {ft}")

    lines.append("")
    lines.append("Moves (full list):")
    for move in moves:
        lines.append(f" - {move}")

    return "\n".join(lines)

async def main():
    pokemon_list = await get_all_pokemon()

    async with aiohttp.ClientSession() as session:
        os.makedirs("pokemon_output", exist_ok=True)

        for p in pokemon_list:
            species = await fetch_json(session, p["species"]["url"])
            evo_json = await fetch_json(session, species["evolution_chain"]["url"])

            evo_chain = extract_evolution_chain(evo_json)
            flavor_texts = extract_english_flavor_texts(species)

            summary = summarize_pokemon(p, species, evo_chain, flavor_texts)

            pokemon_id = p["id"]
            pokemon_name = p["name"].capitalize()

            # TXT guard
            txt_filename = f"{pokemon_id:03d}_{pokemon_name}.txt"
            txt_filepath = os.path.join("pokemon_output", txt_filename)

            if not os.path.exists(txt_filepath):
                with open(txt_filepath, "w", encoding="utf-8") as f:
                    f.write(summary)
                print(f"Created file: {txt_filepath}")
            else:
                print(f"Skipped TXT (already exists): {txt_filepath}")

            # Download images
            front_default_url = p["sprites"].get("front_default")
            artwork_url = p["sprites"].get("other", {}).get("official-artwork", {}).get("front_default")

            front_png = f"{pokemon_id:03d}_{pokemon_name}_front_default.png"
            front_path = os.path.join("pokemon_output", front_png)

            if not os.path.exists(front_path):
                await download_image(session, front_default_url, front_path)
                print(f"Saved sprite: {front_path}")
            else:
                print(f"Skipped sprite (already exists): {front_path}")

            art_png = f"{pokemon_id:03d}_{pokemon_name}_artwork.png"
            art_path = os.path.join("pokemon_output", art_png)

            if not os.path.exists(art_path):
                await download_image(session, artwork_url, art_path)
                print(f"Saved artwork: {art_path}")
            else:
                print(f"Skipped artwork (already exists): {art_path}")

            # DOCX guard
            docx_filename = f"{pokemon_id:03d}_{pokemon_name}.docx"
            docx_filepath = os.path.join("pokemon_output", docx_filename)

            if not os.path.exists(docx_filepath):
                build_docx(pokemon_id, pokemon_name, summary, front_path, art_path)
                print(f"Created DOCX: {docx_filepath}")
            else:
                print(f"Skipped DOCX (already exists): {docx_filepath}")

asyncio.run(main())
