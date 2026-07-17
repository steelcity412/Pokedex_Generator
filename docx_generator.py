import os
import shutil
from docx import Document
from docx.shared import Inches

# ============================================================
# Backup helper
# ============================================================

def backup_file_if_exists(src_path, backup_dir, logger):
    """
    Moves old DOCX files into a backup folder before overwriting.
    This preserves historical versions of each Pokémon document.
    """
    if os.path.exists(src_path):
        os.makedirs(backup_dir, exist_ok=True)
        base_name = os.path.basename(src_path)
        backup_path = os.path.join(backup_dir, base_name)
        shutil.move(src_path, backup_path)
        logger.info(f"Backed up old DOCX to: {backup_path}")


# ============================================================
# DOCX builder
# ============================================================

def build_docx(summary, front_path, art_path, output_path, logger):
    """
    Builds a DOCX file containing:
      - front_default sprite
      - official artwork
      - full summary text (English, Japanese, or Kana)
    """
    try:
        doc = Document()

        # Use a table to align images nicely
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Add front_default sprite if available
        if os.path.exists(front_path):
            cell.add_paragraph().add_run().add_picture(front_path, width=Inches(1.5))

        # Add official artwork if available
        if os.path.exists(art_path):
            cell.add_paragraph().add_run().add_picture(art_path, width=Inches(1.5))

        # Add summary text
        doc.add_paragraph("")
        doc.add_paragraph(summary)

        # Save DOCX
        doc.save(output_path)
        logger.info(f"Saved DOCX: {output_path}")

    except Exception as e:
        logger.error(f"Failed to build DOCX {output_path}: {e}")


# ============================================================
# Main DOCX generator
# ============================================================

def generate_docx_for_lang(pokemon_list, species_map, evo_map, lang, folders, logger):
    """
    Generates DOCX files for a specific language.
    DOCX regeneration only happens if the corresponding TXT file changed.
    (TXT validation + backup happens in txt_generator.py)
    """

    # Determine folders based on language
    if lang == "en":
        docx_dir = folders["docx_en"]
        backup_dir = folders["backups_docx_en"]
        txt_dir = folders["text_en"]
        suffix = "_en"
    elif lang == "ja":
        docx_dir = folders["docx_ja"]
        backup_dir = folders["backups_docx_ja"]
        txt_dir = folders["text_ja"]
        suffix = "_ja"
    elif lang == "ja-Hrkt":
        docx_dir = folders["docx_ja_hrkt"]
        backup_dir = folders["backups_docx_ja_hrkt"]
        txt_dir = folders["text_ja_hrkt"]
        suffix = "_ja-Hrkt"
    else:
        logger.error(f"Unsupported language for DOCX generation: {lang}")
        return

    # Sprite folder (same for all languages)
    sprites_dir = folders["sprites"]

    # Loop through all Pokémon
    for p in pokemon_list:
        pid = p["id"]
        name = p["name"].capitalize()
        id_str = f"{pid:04d}"

        # TXT file must exist first — DOCX depends on TXT
        txt_filename = f"{id_str}_{name}{suffix}.txt"
        txt_path = os.path.join(txt_dir, txt_filename)

        if not os.path.exists(txt_path):
            logger.warning(f"Skipping DOCX ({lang}) — TXT missing: {txt_path}")
            continue

        # Load summary text from TXT
        with open(txt_path, "r", encoding="utf-8") as f:
            summary = f.read()

        # Sprite paths
        front_png = f"{id_str}_{name}_front_default.png"
        art_png = f"{id_str}_{name}_artwork.png"

        front_path = os.path.join(sprites_dir, front_png)
        art_path = os.path.join(sprites_dir, art_png)

        # Output DOCX path
        docx_filename = f"{id_str}_{name}{suffix}.docx"
        docx_path = os.path.join(docx_dir, docx_filename)

        # Backup old DOCX if it exists
        backup_file_if_exists(docx_path, backup_dir, logger)

        # Build new DOCX
        build_docx(summary, front_path, art_path, docx_path, logger)
